"""
Small FastAPI app that serves both the frontend and the chat API.

This keeps development simple:
one command starts the site and the RAG backend together.
"""

from __future__ import annotations

import os
import secrets
import shutil
from pathlib import Path
from typing import Literal

from dotenv import load_dotenv

import pdfplumber
from docx import Document as DocxDocument
from fastapi import Depends, FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, JSONResponse
from fastapi.security import HTTPBasic, HTTPBasicCredentials
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

from backend.rag_service import CHROMA_SQLITE_FILE, ENV_PATH, RAGService, RAGServiceError

load_dotenv()

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FRONTEND_DIR = PROJECT_ROOT / "frontend"
INDEX_FILE = FRONTEND_DIR / "index.html"

app = FastAPI(title="Global Desk")
rag_service = RAGService()
security = HTTPBasic()

class ChatHistoryItem(BaseModel):
    """One prior chat turn sent from the browser."""
    role: Literal["user", "assistant"]
    content: str = Field(min_length=1)


class ChatRequest(BaseModel):
    """Payload for the simple single-chat API."""
    message: str = Field(min_length=1)
    history: list[ChatHistoryItem] = Field(default_factory=list)


def verify_admin(credentials: HTTPBasicCredentials = Depends(security)):
    # compare_digest does a constant-time comparison to prevent timing attacks
    # that could otherwise reveal whether the username or password was correct.
    correct_username = secrets.compare_digest(
        credentials.username, os.getenv("ADMIN_USERNAME", "admin")
    )
    correct_password = secrets.compare_digest(
        credentials.password, os.getenv("ADMIN_PASSWORD", "changeme")
    )
    if not (correct_username and correct_password):
        raise HTTPException(
            status_code=401,
            detail="Incorrect credentials",
            headers={"WWW-Authenticate": "Basic"},
        )


@app.get("/api/health")
def health_check() -> dict[str, bool]:
    return {
        "ok": True,
        "api_key_configured": bool(os.getenv("OPENAI_API_KEY")),
        "knowledge_base_ready": CHROMA_SQLITE_FILE.exists(),
        "env_file_found": ENV_PATH.exists(),
    }


@app.get("/api/admin/verify")
def verify(credentials: HTTPBasicCredentials = Depends(verify_admin)):
    return {"ok": True}


@app.post("/api/admin/upload")
async def upload_document(
    file: UploadFile = File(...),
    credentials: HTTPBasicCredentials = Depends(verify_admin)
):
    """
    Accept a PDF or DOCX upload, extract its text, save it to backend/data/,
    and incrementally add it to the existing ChromaDB knowledge base.
    """
    allowed_types = [
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]

    if file.content_type not in allowed_types:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported.")

    data_dir = PROJECT_ROOT / "backend" / "data"
    data_dir.mkdir(exist_ok=True)

    suffix = ".pdf" if file.content_type == "application/pdf" else ".docx"
    temp_path = data_dir / f"_temp_{file.filename}"
    with open(temp_path, "wb") as f:
        shutil.copyfileobj(file.file, f)

    try:
        if suffix == ".pdf":
            with pdfplumber.open(temp_path) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        else:
            doc = DocxDocument(temp_path)
            text = "\n".join(p.text for p in doc.paragraphs)
    finally:
        temp_path.unlink(missing_ok=True)

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file.")

    txt_filename = Path(file.filename).stem + ".txt"
    txt_path = data_dir / txt_filename
    with open(txt_path, "w") as f:
        f.write(text)

    # Deferred imports — these pull in heavy ML libs and we only need them during upload.
    from backend.scripts.ingestion_pipeline import split_documents
    from langchain_chroma import Chroma
    from langchain_core.documents import Document as LangchainDocument
    from langchain_openai import OpenAIEmbeddings

    new_doc = LangchainDocument(page_content=text, metadata={"source": txt_filename})
    chunks = split_documents([new_doc])

    embedding_model = OpenAIEmbeddings(model="text-embedding-3-small")
    chroma_dir = PROJECT_ROOT / "backend" / "chroma_db"

    vector_db = Chroma(
        persist_directory=str(chroma_dir),
        embedding_function=embedding_model,
        collection_metadata={"hnsw:space": "cosine"}
    )
    vector_db.add_documents(chunks)

    return {"ok": True, "saved_as": txt_filename}


@app.post("/api/chat")
def chat(request: ChatRequest):
    """
    Receive one user message plus prior history and return the RAG answer.
    """
    try:
        return rag_service.answer_question(
            message=request.message,
            history=[item.model_dump() for item in request.history],
        )
    except RAGServiceError as error:
        return JSONResponse(status_code=400, content={"error": str(error)})
    except Exception as error:
        print(f"Unexpected chat error: {error}")
        return JSONResponse(
            status_code=500,
            content={
                "error": (
                    "Something went wrong while generating a response. "
                    "Please try again."
                )
            },
        )


@app.get("/", include_in_schema=False)
def read_index() -> FileResponse:
    return FileResponse(INDEX_FILE)



app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")
app.mount("/admin", StaticFiles(directory=PROJECT_ROOT / "admin", html=True), name="admin")

@app.get("/api/admin/files")
def list_files(credentials: HTTPBasicCredentials = Depends(verify_admin)):
    data_dir = PROJECT_ROOT / "backend" / "data"
    files = [f.name for f in data_dir.glob("*.txt") if f.name != "preview.txt"]
    return {"files": files}


@app.delete("/api/admin/files/{filename}")
def delete_file(filename: str, credentials: HTTPBasicCredentials = Depends(verify_admin)):
    data_dir = PROJECT_ROOT / "backend" / "data"
    file_path = data_dir / filename
    if not file_path.resolve().is_relative_to(data_dir.resolve()):
        raise HTTPException(status_code=400, detail="Invalid filename.")
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found.")
    file_path.unlink()
    return {"ok": True, "deleted": filename}


@app.post("/api/admin/reindex")
def reindex(credentials: HTTPBasicCredentials = Depends(verify_admin)):
    import shutil
    from backend.scripts.ingestion_pipeline import (
        load_urls, load_documents, load_text_files,
        split_documents, vectorize_db
    )
    chroma_dir = PROJECT_ROOT / "backend" / "chroma_db"
    if chroma_dir.exists():
        shutil.rmtree(chroma_dir)
    urls = load_urls()
    url_docs = load_documents(urls)
    text_docs = load_text_files()
    all_docs = url_docs + text_docs
    chunks = split_documents(all_docs)
    vectorize_db(chunks)
    return {"ok": True, "message": "Knowledge base rebuilt successfully."}
